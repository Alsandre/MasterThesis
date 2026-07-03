#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Assemble the GTU-formatted Georgian thesis .docx.

Sources:
  - template:  შაბლონი სამაგისტრო (1).docx   (styles, margins, front-matter pages, title-page text boxes)
  - georgian:  project/draft-en ka.docx        (translated body, flattened)
  - structure: project/draft-en.md             (heading levels + table geometry, via blueprint.py)
  - biblio:    project/draft.md                 ([1]-[35])

Output: project/thesis-ka.docx
"""
import re, copy, sys, zipfile, shutil, os, subprocess, json, io
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from blueprint import parse_blueprint

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, '..', '..'))
def R(p): return os.path.join(ROOT, p)

TEMPLATE = R("შაბლონი სამაგისტრო (1).docx")
KA       = R("project/draft-en ka.docx")
MD       = R("project/draft-en.md")
BIB      = R("project/draft.md")
OUT      = R("project/thesis-ka.docx")

# ---------------- FACTS (confirmed vs placeholder) ----------------
TITLE  = "სასაუბრო ხელოვნური ინტელექტი: ასისტენტიდან თანამოსაუბრემდე ევოლუცია"
AUTHOR = "ალექსანდრე იმნაიშვილი"         # confirmed by Lekso
YEAR   = "2026"                          # cover year (default; editable)
MONTH  = "ივლისი"                        # cover month (default; editable)
SHIFRI = "[შიფრი]"                       # TODO from Lekso
PROGRAM= "[სამაგისტრო პროგრამა]"         # TODO from Lekso
SUPERVISOR = "ოთარ თავდიშვილი"           # confirmed by Lekso
DEFENSE_DATE = "[სხდომის თარიღი]"        # TODO

# chapters rendered cohesively: subsection headings become bold run-in lead-ins
# (flowing prose, no numbered sub-headings) — like the passed GTU reference.
COHESIVE_CHAPTERS = {'1'}

# table captions (numbers MUST match the "ცხრილი N.M" references in the prose)
TABLE_META = [
    ("2.1", "ადამიანის მეხსიერების თეორიები და მათი შესაბამისი LLM-იმპლემენტაციები"),
    ("3.1", "აღრევის ფაქტორების კონტროლი ექსპერიმენტულ დიზაინში"),
    ("4.1", "კონცეფციის დამტკიცების ერთგულება 3.2-ის სრული არქიტექტურის მიმართ"),
    ("4.2", "სესიებს-შორისი გახსენების სიზუსტე პრობის ტიპის მიხედვით"),
    ("4.3", "სიმულირებული აღქმის საშუალო შეფასებები კონსტრუქტების მიხედვით"),
]

# ---------------- proofing corrections (from coherence review) ----------------
CORR_PATH = os.path.join(HERE, 'corrections.json')
CORRECTIONS = json.load(open(CORR_PATH, encoding='utf-8')) if os.path.exists(CORR_PATH) else []

def apply_corrections(t):
    for c in CORRECTIONS:
        t = t.replace(c['old'], c['new'])
    return t

# ---------------- 1. Build poured structure (validated 1:1) ----------------
def pour():
    blocks = parse_blueprint(MD)
    doc = Document(KA)
    ka = [apply_corrections(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    ka_body = ka[1:]                      # [0] is the (re-translated) title -> discard
    out, j = [], 0
    for b in blocks:
        if b[0] in ('h1', 'h2', 'p', 'bullet', 'num'):
            out.append((b[0], ka_body[j])); j += 1
        elif b[0] == 'table':
            grid = []
            for row in b[1]:
                grow = []
                for cell in row:
                    if cell.strip():
                        grow.append(ka_body[j]); j += 1
                    else:
                        grow.append('')
                grid.append(grow)
            out.append(('table', grid))
    assert j == len(ka_body), f"alignment off: consumed {j}/{len(ka_body)}"
    return out

# ---------------- English abstract (for Abstract slot) ----------------
def english_abstract():
    lines = open(MD, encoding='utf-8').read().split('\n')
    s = next(i for i, l in enumerate(lines) if l.strip() == '## Abstract')
    e = next(i for i, l in enumerate(lines) if l.strip().startswith('## §1'))
    paras = []
    for l in lines[s+1:e]:
        t = l.strip()
        if not t or t == '---':
            continue
        t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)   # strip bold
        t = re.sub(r'\*(.+?)\*', r'\1', t)
        paras.append(t)
    return paras   # 5 body paras + keywords line

# ---------------- helpers ----------------
def set_sylfaen(run, size=12, bold=False):
    run.font.name = 'Sylfaen'
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'):
        rf.set(qn(a), 'Sylfaen')

def clear_highlight(run):
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        for h in rpr.findall(qn('w:highlight')):
            rpr.remove(h)

def replace_in_paragraph(p, repls):
    """Replace substrings across a paragraph's concatenated runs.
    Rewrites into a single run using the first run's formatting. For boilerplate lines only."""
    full = ''.join(r.text for r in p.runs)
    new = full
    for a, b in repls:
        new = new.replace(a, b)
    if new == full:
        return False
    if p.runs:
        p.runs[0].text = new
        clear_highlight(p.runs[0])
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(new)
    return True

def para_index(doc, needle):
    """Exact-match a standalone heading/anchor paragraph (avoids substring collisions)."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == needle:
            return i
    for i, p in enumerate(doc.paragraphs):   # fallback: contains
        if needle in p.text:
            return i
    return -1

def insert_paragraph_after(paragraph, text='', style=None, spacing=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, paragraph._parent)
    if style:
        np.style = style
    if text:
        run = np.add_run(text)
        set_sylfaen(run, 12)
    if spacing:
        np.paragraph_format.line_spacing_rule = spacing
    return np

def _insert_in_tblpr(tblPr, el, after_tags):
    """Insert el into tblPr in schema order: before the first of after_tags present."""
    for tag in after_tags:
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            anchor.addprevious(el)
            return
    tblPr.append(el)

def set_table_borders(tbl):
    tblPr = tbl._element.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        borders.append(el)
    # schema order: tblBorders comes before shd/tblLayout/tblLook
    _insert_in_tblpr(tblPr, borders, ('w:shd', 'w:tblLayout', 'w:tblLook'))

def set_col_widths(tbl, grid, total_cm=14.7):
    """Fixed layout with content-proportional column widths (so numeric columns
    stay narrow and long-label columns get the room they need)."""
    ncol = max(len(r) for r in grid)
    tbl.autofit = False        # adds a schema-ordered <w:tblLayout w:type="fixed"/>
    # per-column weight = max content length (floored), dampened for very long cells
    weights = []
    for c in range(ncol):
        maxlen = max((len(r[c]) for r in grid if c < len(r) and r[c]), default=1)
        weights.append(max(4, min(maxlen, 46)))
    s = sum(weights)
    widths = [max(1.3, total_cm * w / s) for w in weights]
    scale = total_cm / sum(widths)             # renormalise to total after the 1.3cm floor
    widths = [w * scale for w in widths]
    twips = [int(round(w * 567)) for w in widths]   # 1cm = 567 twips (dxa)
    # 1) update the tblGrid gridCol widths (authoritative under fixed layout)
    grid_el = tbl._element.find(qn('w:tblGrid'))
    cols = grid_el.findall(qn('w:gridCol'))
    for i, gc in enumerate(cols):
        if i < ncol:
            gc.set(qn('w:w'), str(twips[i]))
    # 2) also set per-cell widths to match
    for row in tbl.rows:
        for i in range(ncol):
            row.cells[i].width = Cm(widths[i])

def clean_prose(t):
    """Light polish: drop the § section-marker so prose matches the §-less headings."""
    return t.replace('§', '')

def rebuild_title_page(doc):
    """Replace the template's floating-box title page (effectively a full-page
    image) with clean inline text + a small centred logo — reproduces the
    f187836 format-session layout so it survives every content rebuild and stays
    searchable / plagiarism-safe."""
    with zipfile.ZipFile(TEMPLATE) as z:
        logo_bytes = z.read('word/media/image1.png')
    # locate the section-0 break paragraph (holds the title-page sectPr)
    sect_p = None
    for p in doc.paragraphs:
        ppr = p._p.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:sectPr')) is not None:
            sect_p = p._p; break
    # delete every paragraph before that break (the floating-box title page)
    for p in list(doc.paragraphs):
        if p._p is sect_p:
            break
        p._p.getparent().remove(p._p)
    # section-0 margins: match the body (38/25) instead of the 5 mm full-bleed
    s0 = doc.sections[0]
    s0.left_margin, s0.right_margin = Mm(38), Mm(25)
    s0.top_margin, s0.bottom_margin = Mm(25), Mm(25)
    ALIGN = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT,
             'left': WD_ALIGN_PARAGRAPH.LEFT}

    def line(text='', img=False, size=12, bold=False, align='center', width_cm=11):
        par = doc.add_paragraph()                 # append (has image-part access)
        par.alignment = ALIGN[align]
        par.paragraph_format.line_spacing = 1.0
        par.paragraph_format.space_after = Pt(0)
        if img:
            par.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(width_cm))
        elif text:
            set_sylfaen(par.add_run(text), size, bold=bold)
        sect_p.addprevious(par._p)                # relocate before the break, in order
        return par

    shifri = SHIFRI if SHIFRI != '[შიფრი]' else '[___]'
    line(img=True, width_cm=11)                   # logo
    line()
    line(AUTHOR, size=12, bold=True, align='right')
    line()
    line(TITLE, size=16, bold=True, align='center')
    line()
    line('შიფრი: ' + shifri, size=14, bold=True, align='center')
    line(); line()
    line('საქართველოს ტექნიკური უნივერსიტეტი', align='center')
    line('თბილისი, 0160, საქართველო', align='center')
    line(f'{MONTH}, {YEAR} წელი', align='center')

def suppress_numbering(p):
    """Override the Heading style's automatic outline numbering (numId=30) with
    numId=0 so only our manual '1.', '1.1' numbers appear (no double-numbering)."""
    ppr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '0')
    numPr.append(numId)
    ppr.append(numPr)

def add_toc_field(paragraph, switches=r'\o "1-2" \h \z \u'):
    """Insert a live Word TOC field. Populates automatically on open in Word
    (updateFields is set); until then it shows the neutral placeholder below."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    fldBegin.set(qn('w:dirty'), 'true')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' TOC {switches} '
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    hint = OxmlElement('w:t'); hint.set(qn('xml:space'), 'preserve')
    hint.text = 'შინაარსი გენერირდება ავტომატურად Word-ში (საჭიროების შემთხვევაში: მარჯვენა ღილაკი → Update Field / F9).'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    r = run._element
    r.append(fldBegin); r.append(instr); r.append(fldSep); r.append(hint); r.append(fldEnd)

# ---------------- main ----------------
def main():
    poured = pour()
    en_abs = english_abstract()
    doc = Document(TEMPLATE)

    # GTU: body 1.5 line spacing (template default is ~single). Set on Normal style.
    nf = doc.styles['Normal'].paragraph_format
    nf.line_spacing = 1.5
    nf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ---- A. front-matter paragraph fills ----
    for p in doc.paragraphs:
        replace_in_paragraph(p, [
            ('სახელი, გვარი', AUTHOR),
            ('სახელი გვარი', AUTHOR),
            ('„------------------“', f'„{TITLE}“'),
            ('„------------------"', f'„{TITLE}“'),
            ('2023 წელი', f'{YEAR} წელი'),
            ('2023 წელს', f'{YEAR} წელს'),
            ('ხელმძღვანელი: _______________', f'ხელმძღვანელი: {SUPERVISOR}'),
        ])
    # commission page: "სამაგისტრო ნაშრომი:  დასახელება"
    for p in doc.paragraphs:
        if 'დასახელება' in p.text and 'სამაგისტრო ნაშრომი' in p.text:
            replace_in_paragraph(p, [('დასახელება', f'„{TITLE}“')])
    # front matter (only template paras exist now): match the GTU sample —
    # (a) left-align justified body paragraphs (justification stretches Georgian),
    # (b) 1.5 line spacing (template used single; sample + GTU body use 1.5).
    for p in doc.paragraphs:
        if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.paragraph_format.line_spacing == 1.0:
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ---- A2. title page: floating boxes -> inline text + small logo ----
    rebuild_title_page(doc)

    # ---- B. რეზიუმე (Georgian abstract) ----
    idx = para_index(doc, 'რეზიუმე')
    anchor = doc.paragraphs[idx]
    # georgian abstract = poured blocks between the 'ანოტაცია' h1 and the next h1 (§1)
    ka_abs, collecting = [], False
    for b in poured:
        if b[0] == 'h1':
            if b[1].strip() == 'ანოტაცია':
                collecting = True; continue
            if collecting:
                break
        elif collecting:
            ka_abs.append(b[1])
    prev = anchor
    for t in ka_abs:
        prev = insert_paragraph_after(prev, clean_prose(t), style='Normal', spacing=WD_LINE_SPACING.SINGLE)

    # ---- C. Abstract (English) ----
    idx = para_index(doc, 'Abstract')
    anchor = doc.paragraphs[idx]
    prev = anchor
    for t in en_abs:
        prev = insert_paragraph_after(prev, t, style='Normal', spacing=WD_LINE_SPACING.SINGLE)

    # ---- D. შინაარსი -> live Word TOC field (auto-populates on open in Word) ----
    idx = para_index(doc, 'შინაარსი')
    toc_p = insert_paragraph_after(doc.paragraphs[idx], '', style='Normal')
    add_toc_field(toc_p)

    # ---- E. body (skip Abstract block; it lives in front matter) ----
    def norm_heading(t):
        t = re.sub(r'^§\s*', '', t)          # drop § from "§1."
        return t
    started = False
    tcount = 0
    current_chapter = None
    pending_leadin = None      # cohesive-chapter run-in label awaiting its paragraph

    def flush_leadin():         # emit a stranded lead-in as its own bold line (rare)
        nonlocal pending_leadin
        if pending_leadin:
            pp = doc.add_paragraph(style='Normal')
            r = pp.add_run(pending_leadin + '.'); set_sylfaen(r, 12, bold=True)
            pending_leadin = None

    for b in poured:
        if b[0] == 'h1':
            if b[1].strip() == 'ანოტაცია':
                continue
            started = True
            m = re.match(r'^(\d+)\.', norm_heading(b[1]))
            current_chapter = m.group(1) if m else None
            p = doc.add_paragraph(style='Heading 1')
            suppress_numbering(p)
            p.paragraph_format.page_break_before = True     # GTU: chapters start on a new page
            run = p.add_run(norm_heading(b[1])); set_sylfaen(run, 14, bold=True)
            continue
        if not started:
            continue
        if b[0] == 'h2':
            if current_chapter in COHESIVE_CHAPTERS:
                pending_leadin = re.sub(r'^\d+\.\d+\s*', '', b[1]).strip()   # drop "1.1 "
                continue
            p = doc.add_paragraph(style='Heading 2')
            suppress_numbering(p)
            run = p.add_run(b[1]); set_sylfaen(run, 13, bold=True)
        elif b[0] == 'p':
            p = doc.add_paragraph(style='Normal')
            if pending_leadin:
                lead = p.add_run(pending_leadin + '. '); set_sylfaen(lead, 12, bold=True)
                pending_leadin = None
            run = p.add_run(clean_prose(b[1])); set_sylfaen(run, 12)
        elif b[0] == 'bullet':
            flush_leadin()
            t = clean_prose(re.sub(r'^•\s*\t?\s*', '', b[1]))
            p = doc.add_paragraph(style='Normal')
            run = p.add_run('•  ' + t); set_sylfaen(run, 12)
        elif b[0] == 'num':
            flush_leadin()
            t = clean_prose(re.sub(r'^(\d+)\s*\t\s*', r'\1. ', b[1]))
            p = doc.add_paragraph(style='Normal'); run = p.add_run(t); set_sylfaen(run, 12)
        elif b[0] == 'table':
            flush_leadin()
            grid = b[1]
            num, title = TABLE_META[tcount]; tcount += 1
            cap = doc.add_paragraph(style='Normal')      # caption above table (GTU)
            cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            crun = cap.add_run(f'ცხრილი {num}. {title}'); set_sylfaen(crun, 11, bold=True)
            ncol = max(len(r) for r in grid)
            tbl = doc.add_table(rows=len(grid), cols=ncol)
            set_table_borders(tbl)
            set_col_widths(tbl, grid)
            for ri, row in enumerate(grid):
                for ci in range(ncol):
                    cell = tbl.cell(ri, ci)
                    txt = clean_prose(row[ci]) if ci < len(row) else ''
                    cp = cell.paragraphs[0]
                    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    run = cp.add_run(txt); set_sylfaen(run, 11, bold=(ri == 0))
            doc.add_paragraph('', style='Normal')

    # ---- F. bibliography ----
    p = doc.add_paragraph(style='Heading 1'); suppress_numbering(p)
    p.paragraph_format.page_break_before = True
    run = p.add_run('ბიბლიოგრაფია'); set_sylfaen(run, 14, bold=True)
    for entry in bibliography():
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(entry); set_sylfaen(run, 11)

    # ---- H. populate ცხრილების ნუსხა (list of tables) ----
    idx = para_index(doc, 'ცხრილების ნუსხა')
    if idx >= 0:
        prev = doc.paragraphs[idx]
        for num, title in TABLE_META:
            prev = insert_paragraph_after(prev, f'ცხრილი {num}. {title}', style='Normal',
                                          spacing=WD_LINE_SPACING.SINGLE)

    # ---- I. cleanup: stray backtick + blank spacer page ----
    # NB: title-page paragraph [0] also anchors floating drawings (logo/title/author
    # text boxes). Only edit the run that holds the backtick — never the drawing runs,
    # since run.text = '' would destroy the drawing.
    for p in doc.paragraphs[:3]:
        for r in p.runs:
            if '`' in r.text:
                r.text = r.text.replace('`', '')
    # delete the run of empty template spacer paragraphs after the English keywords line
    kw = next((i for i, p in enumerate(doc.paragraphs)
               if p.text.strip().startswith('Keywords:')), -1)
    if kw >= 0:
        to_del = []
        for p in doc.paragraphs[kw+1:]:
            if p.text.strip():
                break
            if 'w:br' in p._p.xml and 'type="page"' in p._p.xml:
                break            # keep the page-break paragraph
            to_del.append(p._p)
        for el in to_del:
            el.getparent().remove(el)

    doc.save(OUT)
    print('saved', OUT)

    # ---- G. title-page text boxes (XML surgery, scoped to txbxContent) ----
    patch_textboxes(OUT, {
        'სახელი და გვარი': AUTHOR,
        'ივლისი, 2023 წელი': f'{MONTH}, {YEAR} წელი',
    }, TITLE)
    print('patched text boxes')


def bibliography():
    lines = open(BIB, encoding='utf-8').read().split('\n')
    # scan ONLY within the "## ბიბლიოგრაფია" section
    start = next(i for i, l in enumerate(lines) if l.strip().startswith('## ბიბლიოგრაფია'))
    out, expect = [], 1
    for l in lines[start+1:]:
        s = l.strip()
        if s.startswith('## '):          # next section -> stop
            break
        m = re.match(r'^(\d+)\.\s+(.*)$', s)
        if m and int(m.group(1)) == expect:
            entry = re.sub(r'\s*⚠️.*$', '', m.group(2))
            out.append(f'[{m.group(1)}] {entry}')
            expect += 1
    return out


def patch_textboxes(path, repls, title):
    tmp = path + '.tmp'
    with zipfile.ZipFile(path) as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    xml = data['word/document.xml'].decode('utf-8')

    def fix_txbx(m):
        block = m.group(0)
        for a, b in repls.items():
            block = block.replace(a, b)
        # title placeholder is split across two runs: "ნაშრომის " + "სათაური "
        block = block.replace('ნაშრომის ', title + ' ')
        block = block.replace('სათაური ', '')
        return block
    xml = re.sub(r'<w:txbxContent>.*?</w:txbxContent>', fix_txbx, xml, flags=re.S)
    # NB: the template title page is a fragile floating-box layout (logo image + 4
    # absolutely-positioned text boxes). Scripted resize/reposition of the logo
    # collides the boxes, so the logo size is left as the template designed it —
    # final logo sizing is a trivial manual drag in Word/Pages.
    data['word/document.xml'] = xml.encode('utf-8')

    # tell Word/LibreOffice to update all fields (TOC) on open — insert in schema
    # order (updateFields sits after characterSpacingControl, before footnotePr).
    sx = data['word/settings.xml'].decode('utf-8')
    if 'updateFields' not in sx:
        tag = '<w:updateFields w:val="true"/>'
        for anchor in ('<w:footnotePr', '<w:endnotePr', '<w:compat', '<w:rsids', '</w:settings>'):
            if anchor in sx:
                sx = sx.replace(anchor, tag + anchor, 1)
                break
        data['word/settings.xml'] = sx.encode('utf-8')
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    shutil.move(tmp, path)


if __name__ == '__main__':
    main()
