#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Assemble the ENGLISH-body GTU thesis .docx (parallel to build.py).

Same front matter as the Georgian thesis (title page, recommendation/commission
pages, the Georgian რეზიუმე abstract, the English Abstract, Contents), but the
five body chapters §1–§5 (+ tables) are the ENGLISH text from draft-en.md, and
the reference list is under an English "Bibliography" heading.

Sources:
  - template:  შაბლონი სამაგისტრო (1).docx   (styles, margins, front-matter pages)
  - english:   project/draft-en.md            (body prose + structure, via blueprint.py)
  - georgian:  project/draft-en ka.docx        (only the რეზიუმე abstract is taken from here)
  - biblio:    project/draft.md                 ([1]-[35])

Output: project/thesis-en.docx  (then run: python3 format_thesis.py project/thesis-en.docx)
"""
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

import build
from build import (set_sylfaen, replace_in_paragraph, para_index, insert_paragraph_after,
                   set_table_borders, set_col_widths, clean_prose, rebuild_title_page,
                   fix_front_matter_pagination, suppress_numbering, add_toc_field,
                   english_abstract, patch_textboxes, bibliography, pour,
                   TEMPLATE, TITLE, AUTHOR, YEAR, MONTH, SUPERVISOR, COHESIVE_CHAPTERS)
from blueprint import parse_blueprint, strip_md

MD     = build.MD
OUT_EN = build.R("project/thesis-en.docx")

# English table captions — same numbering as the Georgian TABLE_META (must match
# the "Table N.M" references written into the English prose).
TABLE_META_EN = [
    ("2.1", "Human-memory theories and their corresponding LLM-agent implementations"),
    ("3.1", "Confound control in the experimental design"),
    ("4.1", "Proof-of-concept fidelity against the full architecture of 3.2"),
    ("4.2", "Cross-session recall accuracy by probe type"),
    ("4.3", "Simulated-perception mean ratings by construct"),
]

def norm_heading(t):
    return re.sub(r'^§\s*', '', t)          # drop the § from "§1."


def main():
    poured    = pour()                       # Georgian blocks — only for the რეზიუმე
    en_abs    = english_abstract()           # English Abstract paragraphs + keywords
    en_blocks = parse_blueprint(MD)          # English body (structure + text)
    doc = Document(TEMPLATE)

    # GTU body: 1.5 line spacing on Normal
    nf = doc.styles['Normal'].paragraph_format
    nf.line_spacing = 1.5
    nf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # ---- A. front-matter fills (identical to the Georgian build) ----
    for p in doc.paragraphs:
        replace_in_paragraph(p, [
            ('სახელი, გვარი', AUTHOR),
            ('სახელი გვარი', AUTHOR),
            ('„------------------“', f'„{TITLE}“'),
            ('„------------------"', f'„{TITLE}“'),
            ('2023 წელი', f'{YEAR} წელი'),
            ('2023 წელს', f'{YEAR} წელს'),
            ('ხელმძღვანელი: _______________', f'ხელმძღვანელი: {SUPERVISOR}'),
        ])
    for p in doc.paragraphs:
        if 'დასახელება' in p.text and 'სამაგისტრო ნაშრომი' in p.text:
            replace_in_paragraph(p, [('დასახელება', f'„{TITLE}“')])
    for p in doc.paragraphs:
        if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.paragraph_format.line_spacing == 1.0:
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    rebuild_title_page(doc)                   # Georgian title page (same as thesis-ka)
    fix_front_matter_pagination(doc)

    # ---- B. რეზიუმე (Georgian abstract, from the poured Georgian blocks) ----
    idx = para_index(doc, 'რეზიუმე')
    anchor = doc.paragraphs[idx]
    ka_abs, collecting = [], False
    for b in poured:
        if b[0] == 'h1':
            if b[1].strip() == 'ანოტაცია':
                collecting = True; continue
            if collecting:
                break
        elif collecting:
            ka_abs.append(b[1])
    prev = anchor
    for t in ka_abs:
        prev = insert_paragraph_after(prev, clean_prose(t), style='Normal',
                                      spacing=WD_LINE_SPACING.SINGLE)

    # ---- C. Abstract (English) ----
    idx = para_index(doc, 'Abstract')
    prev = doc.paragraphs[idx]
    for t in en_abs:
        prev = insert_paragraph_after(prev, t, style='Normal',
                                      spacing=WD_LINE_SPACING.SINGLE)

    # ---- D. Contents -> live Word TOC field ----
    idx = para_index(doc, 'შინაარსი')
    toc_p = insert_paragraph_after(doc.paragraphs[idx], '', style='Normal')
    add_toc_field(toc_p)

    # ---- E. body (English) ----
    started = False
    tcount = 0
    current_chapter = None
    pending_leadin = None
    num_counter = 0            # per-list ordinal (blueprint strips the "N." prefix)

    def flush_leadin():
        nonlocal pending_leadin
        if pending_leadin:
            pp = doc.add_paragraph(style='Normal')
            r = pp.add_run(pending_leadin + '.'); set_sylfaen(r, 12, bold=True)
            pending_leadin = None

    for b in en_blocks:
        typ = b[0]
        if typ != 'num':
            num_counter = 0                     # reset ordinal on any non-list block

        if typ == 'h1':
            if b[1].strip() == 'Abstract':
                continue                        # English Abstract lives in front matter
            started = True
            h = norm_heading(b[1])
            m = re.match(r'^(\d+)\.', h)
            current_chapter = m.group(1) if m else None
            p = doc.add_paragraph(style='Heading 1')
            suppress_numbering(p)
            p.paragraph_format.page_break_before = True   # chapters start on a new page
            run = p.add_run(strip_md(h)); set_sylfaen(run, 14, bold=True)
            continue
        if not started:
            continue

        if typ == 'h2':
            if current_chapter in COHESIVE_CHAPTERS:
                pending_leadin = re.sub(r'^\d+\.\d+\s*', '', b[1]).strip()   # drop "1.1 "
                continue
            p = doc.add_paragraph(style='Heading 2')
            suppress_numbering(p)
            run = p.add_run(strip_md(b[1])); set_sylfaen(run, 13, bold=True)
        elif typ == 'p':
            p = doc.add_paragraph(style='Normal')
            if pending_leadin:
                lead = p.add_run(pending_leadin + '. '); set_sylfaen(lead, 12, bold=True)
                pending_leadin = None
            run = p.add_run(strip_md(clean_prose(b[1]))); set_sylfaen(run, 12)
        elif typ == 'bullet':
            flush_leadin()
            t = strip_md(clean_prose(b[1]))
            p = doc.add_paragraph(style='Normal')
            run = p.add_run('•  ' + t); set_sylfaen(run, 12)
        elif typ == 'num':
            flush_leadin()
            num_counter += 1
            t = strip_md(clean_prose(b[1]))
            p = doc.add_paragraph(style='Normal')
            run = p.add_run(f'{num_counter}. ' + t); set_sylfaen(run, 12)
        elif typ == 'table':
            flush_leadin()
            grid = b[1]
            num, title = TABLE_META_EN[tcount]; tcount += 1
            cap = doc.add_paragraph(style='Normal')
            cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            crun = cap.add_run(f'Table {num}. {title}'); set_sylfaen(crun, 11, bold=True)
            ncol = max(len(r) for r in grid)
            tbl = doc.add_table(rows=len(grid), cols=ncol)
            set_table_borders(tbl)
            set_col_widths(tbl, grid)
            for ri, row in enumerate(grid):
                for ci in range(ncol):
                    cell = tbl.cell(ri, ci)
                    txt = strip_md(clean_prose(row[ci])) if ci < len(row) else ''
                    cp = cell.paragraphs[0]
                    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    run = cp.add_run(txt); set_sylfaen(run, 11, bold=(ri == 0))
            doc.add_paragraph('', style='Normal')

    # ---- F. bibliography (English heading; same shared [N] list) ----
    p = doc.add_paragraph(style='Heading 1'); suppress_numbering(p)
    p.paragraph_format.page_break_before = True
    run = p.add_run('Bibliography'); set_sylfaen(run, 14, bold=True)
    for entry in bibliography():
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(entry); set_sylfaen(run, 11)

    # ---- H. populate ცხრილების ნუსხა (list of tables) with English captions ----
    idx = para_index(doc, 'ცხრილების ნუსხა')
    if idx >= 0:
        prev = doc.paragraphs[idx]
        for num, title in TABLE_META_EN:
            prev = insert_paragraph_after(prev, f'Table {num}. {title}', style='Normal',
                                          spacing=WD_LINE_SPACING.SINGLE)

    # ---- I. cleanup: stray backtick + blank spacer pages (same as build.py) ----
    for p in doc.paragraphs[:3]:
        for r in p.runs:
            if '`' in r.text:
                r.text = r.text.replace('`', '')
    kw = next((i for i, p in enumerate(doc.paragraphs)
               if p.text.strip().startswith('Keywords:')), -1)
    if kw >= 0:
        to_del = []
        for p in doc.paragraphs[kw+1:]:
            if p.text.strip():
                break
            if 'w:br' in p._p.xml and 'type="page"' in p._p.xml:
                break
            to_del.append(p._p)
        for el in to_del:
            el.getparent().remove(el)

    def is_break(p):
        return 'w:br' in p._p.xml and 'type="page"' in p._p.xml
    def is_empty(p):
        return not p.text.strip() and '<w:drawing' not in p._p.xml and not is_break(p)
    paras = list(doc.paragraphs)
    to_del = []
    for i, p in enumerate(paras):
        if is_break(p):
            j = i - 1
            while j >= 0 and is_empty(paras[j]):
                to_del.append(paras[j]._p); j -= 1
    for el in to_del:
        if el.getparent() is not None:
            el.getparent().remove(el)

    doc.save(OUT_EN)
    print('saved', OUT_EN)

    patch_textboxes(OUT_EN, {
        'სახელი და გვარი': AUTHOR,
        'ივლისი, 2023 წელი': f'{MONTH}, {YEAR} წელი',
    }, TITLE)
    print('patched text boxes')


if __name__ == '__main__':
    main()
