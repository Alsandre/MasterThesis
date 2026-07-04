#!/usr/bin/env python3
"""
Repeatable, idempotent formatter for the GTU master's thesis.

Run it any time after content changes (new/renamed/reordered headings,
merged sections, etc.). It re-applies all layout rules and REGENERATES the
table of contents from whatever Heading 1/2/3 structure currently exists —
verbatim heading text (so "თავი N." prefixes and cohesive un-subsectioned
chapters are reflected automatically), with dot leaders and real page numbers.

Usage:  python3 format_thesis.py [path/to/thesis.docx]

Deps (already on this machine):
  - python-docx
  - LibreOffice headless  (for pagination -> page numbers)
  - pdftotext (poppler)   (to read the page numbers back)

It does NOT touch wording/content — only formatting + the TOC scaffold.
"""
import sys, os, re, copy, subprocess, tempfile
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DOCX     = sys.argv[1] if len(sys.argv) > 1 else "project/thesis-ka.docx"
SOFFICE  = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
BLACK    = RGBColor(0, 0, 0)
FONT     = "Sylfaen"
H_SIZE   = Pt(14)          # GTU §1.1.6 (reference uses 16 for H1; flip here if desired)
TOC_TAB  = Mm(146.8)       # right dot-leader tab = usable width (38/25 margins)
TOC_INDENT = {1: Mm(0), 2: Mm(8), 3: Mm(16)}
HEAD_STYLES = ('Heading 1', 'Heading 2', 'Heading 3')
TOC_STYLES  = ('TOC 1', 'TOC 2', 'TOC 3')
# (space_before, space_after) in pt — matches reference headings
HEAD_SPACING = {'Heading 1': (18, 4), 'Heading 2': (8, 4), 'Heading 3': (8, 4)}

log = lambda m: print("  -", m)

# script detection for per-run proofing language
GEO_RE = re.compile(r'[Ⴀ-ჿᲐ-Ჿⴀ-⴯]')  # Georgian (Mkhedruli/Mtavruli)
LAT_RE = re.compile(r'[A-Za-z]')
# bibliography identifier patterns -> clickable links
ARXIV_RE = re.compile(r'arXiv:\s?(\d{4}\.\d{4,5})(v\d+)?', re.I)
DOI_RE   = re.compile(r'10\.\d{4,9}/[^\s]+')

# ---------- helpers ----------
def norm(s): return re.sub(r'\s+', ' ', s).strip()

def render(docx, outdir):
    subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx],
                   check=True, capture_output=True)
    return os.path.join(outdir, os.path.splitext(os.path.basename(docx))[0] + ".pdf")

def page_texts(pdf):
    txt = subprocess.run(["pdftotext", "-layout", pdf, "-"],
                         capture_output=True, text=True).stdout
    return txt.split('\f')

def printed_number(page_text):
    n = re.findall(r'(?m)^\s*(\d{1,3})\s*$', page_text)
    return n[-1] if n else None

# ---------- static formatting ----------
def set_margins(d):
    for s in d.sections:
        s.left_margin, s.right_margin = Mm(38), Mm(25)
        s.top_margin, s.bottom_margin = Mm(25), Mm(25)
    log(f"margins 38/25/25/25 on {len(d.sections)} sections")

def format_headings(d):
    for nm in HEAD_STYLES:
        st = d.styles[nm]
        st.font.color.rgb = BLACK; st.font.name = FONT
        st.font.size = H_SIZE; st.font.bold = True
        pf = st.paragraph_format
        pf.alignment = AL.CENTER
        sb, sa = HEAD_SPACING[nm]
        pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    n = 0
    for p in d.paragraphs:
        if p.style.name in HEAD_STYLES and p.text.strip():
            pf = p.paragraph_format
            pf.alignment = AL.CENTER
            sb, sa = HEAD_SPACING[p.style.name]
            pf.space_before = Pt(sb); pf.space_after = Pt(sa)
            for r in p.runs:
                r.font.color.rgb = BLACK; r.font.size = H_SIZE
                r.font.bold = True; r.font.name = FONT
            n += 1
    log(f"headings: black, centered, 14pt bold, ref spacing 18/8+4 ({n} paragraphs)")

def format_abstract(d):
    paras = d.paragraphs
    rez = toc = None
    for i, p in enumerate(paras):
        if p.text.strip() == 'რეზიუმე': rez = i
        if p.style.name == 'TOC Heading' and 'შინაარსი' in p.text: toc = i
    if rez is None or toc is None or rez >= toc:
        log("abstract block not found — skipped"); return
    for i in range(rez, toc):
        p = paras[i]; t = p.text.strip()
        p.paragraph_format.line_spacing = 1.0           # GTU §1.1.8 single
        if t in ('რეზიუმე', 'Abstract'):
            p.paragraph_format.alignment = AL.CENTER
        elif len(t) > 40 and not re.match(r'^(საკვანძო|Keywords)', t):
            p.paragraph_format.alignment = AL.JUSTIFY
    log(f"abstract: single-spaced + justified (paras {rez}..{toc-1})")

def justify_body(d):
    paras = d.paragraphs
    # body starts at first Heading 1 after the TOC heading
    toc_i = next((i for i, p in enumerate(paras)
                  if p.style.name == 'TOC Heading' and 'შინაარსი' in p.text), 0)
    start = next((i for i in range(toc_i + 1, len(paras))
                  if paras[i].style.name == 'Heading 1' and paras[i].text.strip()), None)
    bib = next((i for i, p in enumerate(paras)
                if p.style.name.startswith('Heading') and 'ბიბლიოგრაფია' in p.text), len(paras))
    if start is None:
        log("body start not found — justify skipped"); return
    n = 0
    for i in range(start, bib):
        p = paras[i]; t = p.text.strip()
        if p.style.name == 'Normal' and len(t) > 40 and not t.startswith('[') \
           and not re.match(r'^(ცხრილი|ნახაზი|სურათი)', t):
            p.paragraph_format.alignment = AL.JUSTIFY; n += 1
    log(f"body: justified {n} prose paragraphs (idx {start}..{bib})")

def _set_run_lang(r, code):
    """Set <w:lang w:val> on a run, creating rPr/lang in schema order."""
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); r.insert(0, rPr)
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        # w:lang precedes eastAsianLayout/specVanish/oMath in the rPr schema
        after = next((rPr.find(qn('w:' + t)) for t in ('eastAsianLayout', 'specVanish', 'oMath')
                      if rPr.find(qn('w:' + t)) is not None), None)
        if after is not None: after.addprevious(lang)
        else: rPr.append(lang)
    lang.set(qn('w:val'), code)

def set_language(d, primary='ka-GE', latin='en-US'):
    """Tag EACH run in the language it's actually written in, so Word / Word-for-web
    proof Georgian against the Georgian dictionary and English against English —
    instead of red-squiggling one language as misspellings of the other. This is
    what a well-formed bilingual thesis does (the reference tags every run too).
    A single document default can't serve both scripts. Editor-only cosmetics —
    never affects print/PDF."""
    # document default = Georgian (the thesis's primary language)
    styles = d.styles.element
    for lang in styles.iter(qn('w:lang')):
        lang.set(qn('w:val'), primary)
    # per-run: detect script, tag explicitly (Georgian wins in mixed runs)
    n_ka = n_en = 0
    for r in d.element.body.iter(qn('w:r')):
        txt = ''.join((t.text or '') for t in r.findall(qn('w:t')))
        if not txt.strip():
            continue
        if GEO_RE.search(txt):
            _set_run_lang(r, primary); n_ka += 1
        elif LAT_RE.search(txt):
            _set_run_lang(r, latin); n_en += 1
    log(f"language: default {primary}; per-run tagged {n_ka} {primary} + {n_en} {latin} runs")

def border_tables(d):
    for t in d.tables:
        tblPr = t._tbl.tblPr
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None: tblPr.remove(old)
        b = OxmlElement('w:tblBorders')
        for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement('w:' + e)
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'auto')
            b.append(el)
        # insert in OOXML schema order (tblBorders before shd/tblLayout/tblLook),
        # not at the end — appending after tblLook is invalid and fails validation
        anchor = next((tblPr.find(qn('w:' + tag)) for tag in ('shd', 'tblLayout', 'tblLook')
                       if tblPr.find(qn('w:' + tag)) is not None), None)
        if anchor is not None:
            anchor.addprevious(b)
        else:
            tblPr.append(b)
    log(f"tables: single-line grid borders ({len(d.tables)})")

# ---------- bibliography hyperlinks ----------
def _link_rPr(tmpl):
    """Fresh, schema-ordered rPr for a hyperlink run: blue + underline, inheriting
    font/size/lang from the identifier's original run."""
    rPr = OxmlElement('w:rPr')
    rs = OxmlElement('w:rStyle'); rs.set(qn('w:val'), 'Hyperlink'); rPr.append(rs)
    if tmpl is not None and tmpl.find(qn('w:rFonts')) is not None:
        rPr.append(copy.deepcopy(tmpl.find(qn('w:rFonts'))))
    col = OxmlElement('w:color'); col.set(qn('w:val'), '0563C1'); rPr.append(col)
    if tmpl is not None and tmpl.find(qn('w:sz')) is not None:
        rPr.append(copy.deepcopy(tmpl.find(qn('w:sz'))))
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if tmpl is not None and tmpl.find(qn('w:lang')) is not None:
        rPr.append(copy.deepcopy(tmpl.find(qn('w:lang'))))
    return rPr

def _plain_run(src, text):
    """Clone src run's formatting, replace its text with `text`."""
    r = copy.deepcopy(src)
    for ch in list(r):
        if ch.tag != qn('w:rPr'):
            r.remove(ch)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)
    return r

def _hyperlink(part, url, disp, tmpl):
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), rid)
    r = OxmlElement('w:r'); r.append(_link_rPr(tmpl))
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = disp
    r.append(t); hl.append(r)
    return hl

def _linkify_run(r, part):
    t_el = r.find(qn('w:t'))
    if t_el is None or not t_el.text:
        return 0
    text = t_el.text
    hits = []
    for m in ARXIV_RE.finditer(text):
        hits.append((m.start(), m.end(), m.group(0),
                     'https://arxiv.org/abs/' + m.group(1) + (m.group(2) or '')))
    for m in DOI_RE.finditer(text):
        doi = m.group(0).rstrip('.,;)')
        hits.append((m.start(), m.start() + len(doi), doi, 'https://doi.org/' + doi))
    if not hits:
        return 0
    hits.sort(key=lambda x: x[0])
    pruned, last = [], -1
    for h in hits:                      # drop overlaps (keep leftmost)
        if h[0] >= last:
            pruned.append(h); last = h[1]
    rPr = r.find(qn('w:rPr'))
    nodes, pos = [], 0
    for s, e, disp, url in pruned:
        if s > pos:
            nodes.append(_plain_run(r, text[pos:s]))
        nodes.append(_hyperlink(part, url, disp, rPr))
        pos = e
    if pos < len(text):
        nodes.append(_plain_run(r, text[pos:]))
    for nd in nodes:
        r.addprevious(nd)
    r.getparent().remove(r)
    return len(pruned)

def linkify_bibliography(d):
    """Wrap every arXiv id / DOI in the bibliography in a real clickable hyperlink
    (blue + underline) WITHOUT changing any wording. Idempotent: already-linked
    identifiers live inside <w:hyperlink> and aren't direct <w:r> children, so a
    re-run skips them."""
    part = d.part
    paras = d.paragraphs
    start = next((i for i, p in enumerate(paras)
                  if p.style.name.startswith('Heading')
                  and re.search(r'ბიბლიოგრაფია|ლიტერატურა', p.text)), 0)
    total = 0
    for p in paras[start:]:
        for r in list(p._p.findall(qn('w:r'))):   # direct children only -> idempotent
            total += _linkify_run(r, part)
    log(f"bibliography: linkified {total} identifiers (arXiv + DOI), from para {start}")

# ---------- TOC ----------
def collect_headings(d):
    out = []
    for p in d.paragraphs:
        if p.style.name in HEAD_STYLES and p.text.strip():
            out.append((int(p.style.name[-1]), p.text.strip()))
    return out

STOP_MARKERS = ('სურათების ნუსხა', 'ცხრილების ნუსხა', 'ნახაზების ნუსხა',
                'ილუსტრაციების ნუსხა')

def clear_old_toc(d):
    """Remove ONLY the TOC entries — everything between the 'შინაარსი' heading
    and the next page break / heading / list-of-figures|tables marker.
    Preserves the list-of-figures/tables sections that follow the TOC."""
    paras = d.paragraphs
    toc_i = next((i for i, p in enumerate(paras)
                  if p.style.name == 'TOC Heading' and 'შინაარსი' in p.text), None)
    if toc_i is None:
        return 0
    end = len(paras)
    for i in range(toc_i + 1, len(paras)):
        p = paras[i]; xml = p._p.xml
        pPr = p._p.find(qn('w:pPr'))
        is_break = ('w:br' in xml and 'type="page"' in xml) or \
                   (pPr is not None and pPr.find(qn('w:sectPr')) is not None)
        is_head = p.style.name.startswith('Heading') or p.style.name == 'TOC Heading'
        is_marker = any(m in p.text for m in STOP_MARKERS)
        if is_break or is_head or is_marker:
            end = i; break
    removed = 0
    for p in list(paras[toc_i + 1:end]):
        p._p.getparent().remove(p._p); removed += 1
    return removed

def build_entries(d, headings, page_of):
    toc_h = next((p for p in d.paragraphs
                  if p.style.name == 'TOC Heading' and 'შინაარსი' in p.text), None)
    assert toc_h is not None, "no 'შინაარსი' TOC Heading found"
    anchor = toc_h._p
    for lvl, text in headings:
        p = d.add_paragraph()                    # Normal + direct fmt (Word-robust)
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.left_indent = TOC_INDENT[lvl]
        # dot-leader right tab DIRECTLY on the paragraph — Word renders this
        pf.tab_stops.add_tab_stop(TOC_TAB, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        for chunk in (text, '\t', str(page_of.get(text, '?'))):
            r = p.add_run(chunk); r.font.name = FONT; r.font.size = Pt(12)
        anchor.addnext(p._p); anchor = p._p

def regen_toc(d, docx_path):
    headings = collect_headings(d)
    clear_old_toc(d)
    # pass 1: placeholder pages -> fixes TOC length/pagination
    PLACE = 'XPGX'
    build_entries(d, headings, {t: PLACE for _, t in headings})
    tmp = tempfile.mkdtemp()
    d.save(docx_path)
    pdf = render(docx_path, tmp)
    pages = page_texts(pdf)
    toc_pages = {i for i, pg in enumerate(pages) if PLACE in pg}
    page_of = {}
    for _, text in headings:
        key = norm(text)
        for i, pg in enumerate(pages):
            if i in toc_pages: continue
            if key in norm(pg):
                page_of[text] = printed_number(pages[i]) or '?'; break
        page_of.setdefault(text, '?')
    # pass 2: real pages (same entry count -> pagination stable)
    clear_old_toc(d)
    build_entries(d, headings, page_of)
    missing = [t for t in page_of if page_of[t] == '?']
    log(f"TOC regenerated: {len(headings)} entries; unresolved pages: {missing or 'none'}")

# ---------- main ----------
def main():
    print(f"Formatting {DOCX} ...")
    d = Document(DOCX)
    set_margins(d)
    format_headings(d)
    format_abstract(d)
    justify_body(d)
    border_tables(d)
    set_language(d)
    linkify_bibliography(d)
    regen_toc(d, DOCX)          # saves internally
    d.save(DOCX)
    print("DONE.")

if __name__ == '__main__':
    main()
